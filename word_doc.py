def create_word_doc_title(doc_title = ''):
    from docx import Document
    from docx.shared import Inches, Pt
    doc = Document()
    doc.add_heading(doc_title, 0)
    doc.add_picture('ucs.png', width=Inches(3.25))
    doc.add_page_break()

    return doc


def create_word_doc_paragraph(doc, heading_text = '', heading_level = 1,
                            paragraph_text = ''):
    from docx.enum.section import WD_ORIENT, WD_SECTION
    doc.add_heading(heading_text, level=heading_level)
    #new_section.orientation = WD_ORIENT.LANDSCAPE
    p = doc.add_paragraph(paragraph_text)
    current_section = doc.sections[-1]
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.page_width = new_width
    new_section.page_height = new_height

    return doc


def create_word_doc_table(doc, df):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'Light List Accent 1')

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]



    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    doc.add_page_break()
    return doc
